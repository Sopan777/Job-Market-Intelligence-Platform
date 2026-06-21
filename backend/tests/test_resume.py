import io
import pytest
from unittest.mock import patch

from src.analyzer.resume import (
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_resume_text,
)


def _make_docx_bytes(text: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_minimal_pdf_bytes() -> bytes:
    # Minimal valid PDF with one text object
    return b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj
4 0 obj << /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (Python SQL Docker) Tj ET
endstream
endobj
5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
xref
0 6
0000000000 65535 f\r
0000000009 00000 n\r
0000000058 00000 n\r
0000000115 00000 n\r
0000000274 00000 n\r
0000000370 00000 n\r
trailer << /Size 6 /Root 1 0 R >>
startxref
441
%%EOF"""


class TestExtractTextFromDocx:
    def test_roundtrip(self):
        text = "Python SQL Docker Kubernetes"
        result = extract_text_from_docx(_make_docx_bytes(text))
        assert "Python" in result
        assert "SQL" in result

    def test_empty_bytes_returns_empty(self):
        result = extract_text_from_docx(b"")
        assert result == ""

    def test_invalid_bytes_returns_empty(self):
        result = extract_text_from_docx(b"not a docx file")
        assert result == ""


class TestExtractTextFromPdf:
    def test_returns_string(self):
        result = extract_text_from_pdf(_make_minimal_pdf_bytes())
        assert isinstance(result, str)

    def test_empty_bytes_returns_empty(self):
        result = extract_text_from_pdf(b"")
        assert result == ""

    def test_invalid_bytes_returns_empty(self):
        result = extract_text_from_pdf(b"not a pdf")
        assert result == ""


class TestExtractResumeText:
    def test_dispatches_docx(self):
        text = "TensorFlow PyTorch scikit-learn"
        result = extract_resume_text(_make_docx_bytes(text), "resume.docx")
        assert "TensorFlow" in result

    def test_dispatches_pdf(self):
        with patch("src.analyzer.resume.extract_text_from_pdf", return_value="Python SQL") as mock_pdf:
            result = extract_resume_text(b"fake", "cv.pdf")
        mock_pdf.assert_called_once()
        assert result == "Python SQL"

    def test_unsupported_extension_raises(self):
        with pytest.raises(ValueError, match="Unsupported"):
            extract_resume_text(b"data", "resume.txt")

    def test_no_extension_raises(self):
        with pytest.raises(ValueError):
            extract_resume_text(b"data", "resume")

    def test_empty_docx_returns_empty(self):
        from docx import Document
        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = extract_resume_text(buf.getvalue(), "empty.docx")
        assert result == ""

    def test_whitespace_normalised(self):
        with patch("src.analyzer.resume.extract_text_from_docx", return_value="Python   SQL\n\nDocker"):
            result = extract_resume_text(b"fake", "r.docx")
        assert "  " not in result
        assert "\n" not in result
